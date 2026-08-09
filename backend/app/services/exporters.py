from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile
import fitz
from docx import Document
from markdown import markdown
from sqlalchemy.orm import Session
from ..errors import AppError
from ..models import Course, Lecture
from ..utils.security import sanitize_filename


def lecture_markdown(lecture: Lecture) -> str:
    if not lecture.note:
        raise AppError("RESOURCE_NOT_FOUND", "No notes are available for this lecture.", 404)
    return lecture.note.content_markdown


def markdown_to_docx(title: str, body: str) -> bytes:
    doc = Document()
    doc.add_heading(title, 0)
    for line in body.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], 1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], 2)
        elif line.strip():
            doc.add_paragraph(line)
    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()


def markdown_to_pdf(title: str, body: str) -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    text = f"{title}\n\n{body}"
    page.insert_textbox(fitz.Rect(50, 50, 545, 790), text, fontsize=10)
    return doc.tobytes()


def course_body(course: Course) -> str:
    parts = [f"# {course.title}"]
    course_guide = next((artifact for artifact in course.artifacts if artifact.artifact_type == "course"), None)
    if course_guide:
        parts.append("# Course Revision Guide\n\n" + course_guide.content_markdown)
    week_artifacts = {artifact.week_number: artifact for artifact in course.artifacts if artifact.artifact_type == "week"}
    for week in sorted({lecture.week_number for lecture in course.lectures}):
        parts.append(f"# Week {week}")
        if week in week_artifacts:
            parts.append("## Week Summary\n\n" + week_artifacts[week].content_markdown)
        for lecture in sorted([item for item in course.lectures if item.week_number == week and item.note], key=lambda item: item.lecture_number):
            parts.append(f"## Lecture {lecture.lecture_number}: {lecture.title}\n\n{lecture.note.content_markdown}")
    return "\n\n".join(parts)


def obsidian_zip(course: Course) -> bytes:
    stream = BytesIO()
    root = sanitize_filename(course.title)
    with ZipFile(stream, "w", ZIP_DEFLATED) as zip_file:
        zip_file.writestr(f"{root}/00-Course-Overview.md", f"---\ncourse: {course.title}\ntags:\n  - nptel\n---\n\n# {course.title}\n")
        index = [f"# {course.title} Index"]
        for lecture in course.lectures:
            if not lecture.note:
                continue
            filename = f"{root}/Week-{lecture.week_number:02d}/Lecture-{lecture.lecture_number:02d}-{sanitize_filename(lecture.title)}.md"
            index.append(f"- [[Week-{lecture.week_number:02d}/Lecture-{lecture.lecture_number:02d}-{sanitize_filename(lecture.title)}|{lecture.title}]]")
            zip_file.writestr(filename, f"---\ncourse: {course.title}\nweek: {lecture.week_number}\nlecture: {lecture.lecture_number}\ninstructor: {course.instructor or ''}\nsource: NPTEL\ntranscript_source: {lecture.transcript.source if lecture.transcript else 'unavailable'}\ntags:\n  - nptel\n  - lecture-notes\n---\n\n{lecture.note.content_markdown}")
        zip_file.writestr(f"{root}/01-Course-Index.md", "\n".join(index))
        zip_file.writestr(f"{root}/Glossary.md", "# Glossary\n")
        zip_file.writestr(f"{root}/Practice-Questions.md", "# Practice Questions\n")
    return stream.getvalue()
